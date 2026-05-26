"""Unit tests for cookbook.scars.import_aware_imports."""

from __future__ import annotations

from cookbook.scars.import_aware_imports import (
    DocxImportReminderScar,
    ImportAwareWriteScar,
    imported_top_level_packages,
)
from fscars.core.payload import HookEventType, HookPayload


def _payload(file_path: str, content: str) -> HookPayload:
    return HookPayload(
        event_type=HookEventType.PRE_TOOL_USE,
        tool_name="Write",
        tool_input={"file_path": file_path, "content": content},
        prompt=None,
        cwd="/tmp",
        session_id="s1",
        raw={},
    )


def test_ast_detects_simple_import():
    assert imported_top_level_packages("import docx\n") == {"docx"}


def test_ast_detects_from_import():
    assert imported_top_level_packages("from docx import Document\n") == {"docx"}


def test_ast_detects_dotted_import():
    assert imported_top_level_packages(
        "from docx.shared import Inches\n"
    ) == {"docx"}


def test_ast_ignores_relative_imports():
    code = "from . import sibling\nfrom .submod import x\n"
    assert imported_top_level_packages(code) == set()


def test_ast_strings_mentioning_package_do_not_count():
    code = '"""tutorial about docx"""\nprint("import docx")\n'
    assert imported_top_level_packages(code) == set()


def test_regex_fallback_on_syntax_error():
    broken = "import docx\ndef oops(:\n"  # invalid syntax
    assert "docx" in imported_top_level_packages(broken)


class _Concrete(ImportAwareWriteScar):
    scar_id = "test-scar"
    name = "Test"
    rule = "test rule"
    watched_packages = ("docx",)


def test_matches_when_package_imported():
    scar = _Concrete()
    assert scar.matches(_payload("foo.py", "import docx\n")) is True


def test_does_not_match_for_non_python():
    scar = _Concrete()
    assert scar.matches(_payload("foo.md", "import docx\n")) is False


def test_does_not_match_when_path_excluded():
    scar = _Concrete()
    assert scar.matches(_payload("test_foo.py", "import docx\n")) is False
    assert scar.matches(_payload("/hook_x.py", "import docx\n")) is False
    assert scar.matches(_payload("foo/__pycache__/x.py", "import docx\n")) is False


def test_does_not_match_when_string_only_mention():
    scar = _Concrete()
    assert (
        scar.matches(_payload("foo.py", '"""mentions docx in docstring"""\n'))
        is False
    )


class _PipelineConcrete(ImportAwareWriteScar):
    scar_id = "test-pipeline"
    name = "Test"
    rule = "test rule"
    watched_packages = ("docx",)
    pipeline_path_fragments = ("/reports/",)
    usage_patterns = (r"\bDocument\s*\(",)


def test_matches_via_usage_hint_in_pipeline_path():
    scar = _PipelineConcrete()
    # No import statement, but Document() call inside a pipeline path
    p = _payload(
        "/proj/reports/generate.py",
        "doc = Document()\ndoc.save('out.docx')\n",
    )
    assert scar.matches(p) is True


def test_usage_hint_outside_pipeline_does_not_match():
    scar = _PipelineConcrete()
    p = _payload(
        "/proj/scripts/random.py",
        "doc = Document()\n",
    )
    assert scar.matches(p) is False


def test_trigger_match_names_imports():
    scar = _Concrete()
    trigger = scar.trigger_match(_payload("foo.py", "import docx\n"))
    assert "foo.py" in trigger
    assert "docx" in trigger


def test_docx_example_disabled_by_default():
    assert DocxImportReminderScar.enabled is False


def test_build_output_uses_overrides():
    scar = DocxImportReminderScar()
    out = scar.build_output(_payload("/proj/reports/x.py", "import docx\n"))
    assert "docx-import-reminder" in out.additional_context
    assert "post-process" in out.additional_context
